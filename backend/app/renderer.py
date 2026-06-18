"""Renders structured resume JSON into the Australian-standard .docx.
Python-docx port of the validated template (A4, 9.5pt body, navy headings,
right-aligned dates, ATS-safe single column)."""

import io
from docx import Document
from docx.shared import Pt, RGBColor, Mm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Calibri"
INK = RGBColor(0x1A, 0x1A, 0x1A)
ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
MUTED = RGBColor(0x55, 0x55, 0x55)
RULE = "BFC9D6"
RIGHT_TAB = Twips(9100)  # right edge of content column on A4 w/ these margins


def _run(p, text, *, size=9.5, color=INK, bold=False, italic=False):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    return r


def _space(p, before=0, after=4, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line


def _bottom_border(p, color=RULE, size=6):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def _heading(doc, text):
    p = doc.add_paragraph()
    _space(p, before=6.5, after=2.2)
    r = _run(p, text.upper(), size=10.5, color=ACCENT, bold=True)
    r.font.element.rPr.append(_caps_spacing())
    _bottom_border(p)
    return p


def _caps_spacing():
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:val"), "20")
    return spc


def _right_tab(p):
    p.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)


def _entry(doc, e):
    p = doc.add_paragraph()
    _space(p, before=3, after=0.4)
    _run(p, e.get("org", ""), size=9.5, color=INK, bold=True)
    if e.get("dates"):
        _right_tab(p)
        _run(p, "\t" + e["dates"], size=8.5, color=MUTED)
    if e.get("role") or e.get("location"):
        sub = doc.add_paragraph()
        _space(sub, after=1)
        if e.get("role"):
            _run(sub, e["role"], size=9.5, color=ACCENT, italic=True)
        if e.get("role") and e.get("location"):
            _run(sub, "  \u2014  ", size=8.5, color=MUTED)
        if e.get("location"):
            _run(sub, e["location"], size=8.5, color=MUTED)
    for b in e.get("bullets", []):
        bp = doc.add_paragraph(style="List Bullet")
        _space(bp, after=0.5, line=1.0)
        bp.paragraph_format.left_indent = Pt(13.7)
        bp.paragraph_format.first_line_indent = Pt(-8.7)
        if isinstance(b, dict):  # {lead, text}
            _run(bp, b.get("lead", "") + " ", bold=True)
            _run(bp, b.get("text", ""))
        else:
            _run(bp, str(b))


def build_docx(data: dict) -> bytes:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)  # A4
    sec.top_margin = Twips(680)
    sec.bottom_margin = Twips(480)
    sec.left_margin = Twips(950)
    sec.right_margin = Twips(950)

    base = doc.styles["Normal"]
    base.font.name = FONT
    base.font.size = Pt(9.5)
    base.font.color.rgb = INK

    # Header
    p = doc.add_paragraph(); _space(p, after=1)
    _run(p, data.get("name", ""), size=15, color=ACCENT, bold=True)
    if data.get("tagline"):
        tp = doc.add_paragraph(); _space(tp, after=2)
        _run(tp, data["tagline"], size=9.5, color=MUTED)
    c = data.get("contact", {})
    contact = "  \u00b7  ".join(x for x in
                                [c.get("location"), c.get("phone"),
                                 c.get("email"), c.get("links")] if x)
    cp = doc.add_paragraph(); _space(cp, after=4)
    _run(cp, contact, size=9, color=MUTED)
    _bottom_border(cp, size=4)

    if data.get("profile"):
        _heading(doc, "Professional Profile")
        pr = doc.add_paragraph(); _space(pr, after=1.5, line=1.05)
        _run(pr, data["profile"])

    if data.get("skills"):
        _heading(doc, "Key Skills")
        bare = []
        for s in data["skills"]:
            label = (s.get("label") or "").strip()
            detail = (s.get("detail") or "").strip()
            if detail:
                sp = doc.add_paragraph(); _space(sp, after=1, line=1.0)
                _run(sp, label + ": ", bold=True)
                _run(sp, detail)
            elif label:
                bare.append(label)
        if bare:
            sp = doc.add_paragraph(); _space(sp, after=1, line=1.0)
            _run(sp, "Additional skills: ", bold=True)
            _run(sp, "  \u00b7  ".join(bare))

    if data.get("projects"):
        _heading(doc, "Projects")
        for e in data["projects"]:
            _entry(doc, e)

    if data.get("experience"):
        _heading(doc, "Professional Experience")
        for e in data["experience"]:
            _entry(doc, e)

    if data.get("education"):
        _heading(doc, "Education")
        for ed in data["education"]:
            ep = doc.add_paragraph(); _space(ep, before=2, after=0.3 if ed.get("sub") else 1.2)
            _run(ep, ed.get("left", ""), size=9.5, bold=True)
            if ed.get("dates"):
                _right_tab(ep)
                _run(ep, "\t" + ed["dates"], size=8.5, color=MUTED)
            if ed.get("sub"):
                sp = doc.add_paragraph(); _space(sp, after=1.2)
                _run(sp, ed["sub"], size=8.5, color=MUTED)

    if data.get("certifications"):
        _heading(doc, "Licences & Certifications")
        cp2 = doc.add_paragraph(); _space(cp2, after=1, line=1.0)
        _run(cp2, data["certifications"], size=8.5)

    rp = doc.add_paragraph(); _space(rp, before=3)
    _run(rp, "References: available on request.", size=8.5, color=MUTED, italic=True)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
